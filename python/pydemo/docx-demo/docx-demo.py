from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pypinyin import lazy_pinyin

# 创建一个新的Word文档
doc = Document()

# 设置文档的基本样式
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 定义要标注拼音的中文文本
chinese_text = "你好，世界！"

# 添加段落到文档
paragraph = doc.add_paragraph()

# 遍历中文文本中的每个字符
for char in chinese_text:
    if '\u4e00' <= char <= '\u9fff':  # 判断是否为中文字符
        # 获取拼音
        pinyin = lazy_pinyin(char)[0]

        # 创建 Ruby Text 结构
        ruby_xml = f'''
            <w:r {nsdecls('w')}>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体" />
                    <w:sz w:val="24" />
                </w:rPr>
                <w:ruby>
                    <w:rubyBase>
                        <w:r>
                            <w:t>{char}</w:t>
                        </w:r>
                    </w:rubyBase>
                    <w:rt>
                        <w:r>
                            <w:t>{pinyin}</w:t>
                        </w:r>
                    </w:rt>
                </w:ruby>
            </w:r>
        '''

        # 将 Ruby Text 结构添加到段落
        ruby_element = parse_xml(ruby_xml)
        paragraph._p.append(ruby_element)

        # 添加空格以分隔字符
        paragraph.add_run(" ")
    else:
        paragraph.add_run(char)  # 非中文字符直接添加

# 保存文档
doc.save("chinese_with_pinyin.docx")