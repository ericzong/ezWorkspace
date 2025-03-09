import re
import sys

import pinyin_jyutping
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from opencc import OpenCC


class LrcJyutpingGenerator:
    """
    描述：输入一个文本文件生成一个 Word 文档。其中：
    1. 将每一文本行转换为繁体中文
    2. 为每个繁体中文添加粤拼注音
    """

    def __init__(self):
        self.trans = JyutpingTransliterator()
        self.pattern = r'([\u4e00-\u9fff]+|[^\u4e00-\u9fff]+)' # 匹配汉字和非汉字字符

    def __split(self, jyutping: str) -> list[str]:
        return re.findall(self.pattern, jyutping)

    @staticmethod
    def __is_chinese(char: str) -> bool:
        return '\u4e00' <= char[0] <= '\u9fff'

    def process_file(self, lrc_path: str):
        with open(lrc_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 创建一个新的 Word 文档
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '楷体'
        font.size = Pt(12)

        for line in lines:
            paragraph = doc.add_paragraph()
            # 移除行尾换行符
            line = line.rstrip('\n')
            fragments = self.__split(line)

            for fragment in fragments:
                if self.__is_chinese(fragment):
                    hk_chinese, jyutping = self.trans.transliterate(fragment)
                    for base, rt in zip(hk_chinese, jyutping.split(' ')):  # 粤拼是用空格分隔的
                        DocxLrcXmlWriter(paragraph).write_lrc(base, rt)
                        # 添加空格以分隔字符
                        paragraph.add_run(" ")
                else:
                    run = paragraph.add_run(fragment)
                    paragraph.add_run(" ")

        doc.save(lrc_path + '.docx')


class JyutpingTransliterator:
    def __init__(self):
        # 创建一个简体中文到繁体中文的转换器
        self.cc = OpenCC('s2hk')
        self.j = pinyin_jyutping.PinyinJyutping()

    def transliterate(self, lrc: str) -> tuple[str, str]:
        hk_chinese = self.cc.convert(lrc)
        jyutping = self.j.jyutping(hk_chinese, tone_numbers=True, spaces=True)
        return hk_chinese, jyutping


class DocxLrcXmlWriter:
    def __init__(self, paragraph):
        self.paragraph = paragraph

    def write_lrc(self, base: str, rt: str):
        pre_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="楷体" w:cs="楷体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="zh-CN"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
          </w:r>
        '''

        post_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="楷体" w:cs="楷体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="zh-CN"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
          </w:r>
        '''

        ruby_xml = r'''
          <w:r {0}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="楷体" w:cs="楷体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="zh-CN"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> EQ \* jc0 \* "Font:微软雅黑" \* hps12 \o \ad(\s \up 15({1}),{2})</w:instrText>
          </w:r>
        '''.format(nsdecls('w'), rt, base)

        self.paragraph._p.append(parse_xml(pre_ele))
        self.paragraph._p.append(parse_xml(ruby_xml))
        self.paragraph._p.append(parse_xml(post_ele))


if __name__ == "__main__":
    if len(sys.argv) == 2:
        file_path = sys.argv[1]
    else:
        print("请输入要处理的粤语歌词LRC文件路径：")
        file_path = input()

    LrcJyutpingGenerator().process_file(file_path)
