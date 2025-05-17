import sys

import fugashi
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt


# 初始化 Jamdict 对象
# jam = Jamdict()

# 日语歌词生成器
# 描述：输入一个文本文件生成一个 Word 文档。其中：
#   1. 为每个日本汉字添加平假名注音
class LrcHiraganaGenerator:
    def __init__(self):
        self.trans = JapaneseTransliterator()
        self.default_font = 'BIZ UDMincho Medium'
        # 创建一个新的 Word 文档
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = self.default_font
        font.size = Pt(14)
        self.doc = doc

    def process_file(self, lrc_path: str):
        with open(lrc_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            paragraph = self.doc.add_paragraph()
            line = line.replace(' ', '　')
            for word in self.trans.tagger(line):
                if word.feature.kana:
                    # 将片假名转换为平假名
                    hiragana = self.trans.katakana_to_hiragana(word.feature.kana)
                    # 如果有汉字，标注平假名
                    if word.char_type == 2:  # 2 表示（包含）汉字
                        (base, rt, common_suffix) = self.trans.find_common_suffix(word.surface, hiragana)
                        JpDocxLrcXmlWriter(paragraph, base_font=self.default_font,
                                           ruby_font=self.default_font).write_lrc(base, rt)
                        paragraph.add_run(common_suffix)
                    else:  # 3: 符号；4: 数字；5：字母；6：假名
                        paragraph.add_run(word.surface.replace('　', ' '))
                else:
                    paragraph.add_run(word.surface.replace('　', ' '))

        self.doc.save(lrc_path + '.ja.docx')


class JapaneseTransliterator:
    def __init__(self):
        """
        fugashi 是一个用于日语分词的 Python 库。
        Tagger 类能将输入的日语文本分割成一个个单词（或称“词元”），并提供每个单词的词性、读音等信息。
        """
        self.tagger = fugashi.Tagger()

    def katakana_to_hiragana(self, kana):
        """将片假名转换为平假名"""
        hiragana = ""
        for char in kana:
            if 0x30A0 <= ord(char) <= 0x30FF:  # 检查是否为片假名
                hiragana += chr(ord(char) - 0x60)  # 转换为平假名
            else:
                hiragana += char
        return hiragana

    def find_common_suffix(self, str1, str2):
        """
        找出两个字符串的共同后缀，并返回一个元组。
        元组的第一个元素是第一个字符串去掉共同后缀后的部分，
        第二个元素是第二个字符串去掉共同后缀后的部分，
        第三个元素是两个字符串的共同后缀。

        :param str1: 第一个输入字符串
        :param str2: 第二个输入字符串
        :return: 包含三个元素的元组
        """
        min_len = min(len(str1), len(str2))
        suffix_length = 0

        # 从后往前遍历，找出共同后缀的长度
        for i in range(1, min_len + 1):
            if str1[-i] == str2[-i]:
                suffix_length = i
            else:
                break

        common_suffix = str1[-suffix_length:] if suffix_length > 0 else ""
        remaining_str1 = str1[:-suffix_length] if suffix_length > 0 else str1
        remaining_str2 = str2[:-suffix_length] if suffix_length > 0 else str2

        return remaining_str1, remaining_str2, common_suffix


class DocxLrcXmlWriter:
    def __init__(self, paragraph):
        self.paragraph = paragraph

    def write_lrc(self, base: str, rt: str):
        pre_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="宋体" w:cs="宋体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
          </w:r>
        '''

        post_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="宋体" w:cs="宋体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
          </w:r>
        '''

        ruby_xml = r'''
          <w:r {0}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="宋体" w:cs="宋体"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> EQ \* jc0 \* "Font:微软雅黑" \* hps12 \o \ad(\s \up 15({1}),{2})</w:instrText>
          </w:r>
        '''.format(nsdecls('w'), rt, base)

        self.paragraph._p.append(parse_xml(pre_ele))
        self.paragraph._p.append(parse_xml(ruby_xml))
        self.paragraph._p.append(parse_xml(post_ele))


class JpDocxLrcXmlWriter:
    def __init__(self, paragraph, base_font="宋体", ruby_font="微软雅黑"):
        self.paragraph = paragraph
        self.base_font = base_font
        self.ruby_font = ruby_font

    def write_lrc(self, base: str, rt: str):
        pre_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="{self.base_font}" w:cs="{self.base_font}"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
          </w:r>
        '''

        post_ele = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="{self.base_font}" w:cs="{self.base_font}"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
          </w:r>
        '''

        ruby_xml = f'''
          <w:r {nsdecls('w')}>
            <w:rPr>
              <w:rFonts w:hint="eastAsia" w:eastAsia="{self.base_font}" w:cs="{self.base_font}"/>
              <w:sz w:val="28"/>
              <w:szCs w:val="28"/>
              <w:shd w:val="clear" w:color="auto" w:fill="auto"/>
              <w:lang w:eastAsia="ja-JP"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> EQ \\* jc0 \\* "Font:{self.ruby_font}" \\* hps12 \\o \\ad(\\s \\up 15({rt}),{base})</w:instrText>
          </w:r>
        '''

        self.paragraph._p.append(parse_xml(pre_ele))
        self.paragraph._p.append(parse_xml(ruby_xml))
        self.paragraph._p.append(parse_xml(post_ele))


if __name__ == "__main__":
    if len(sys.argv) == 2:
        file_path = sys.argv[1]
    else:
        print("请输入要处理的日语歌词文件路径：")
        file_path = input()

    LrcHiraganaGenerator().process_file(file_path)
